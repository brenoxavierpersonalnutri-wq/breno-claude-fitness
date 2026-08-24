#!/usr/bin/env python3
"""
gerar_briefing_docx.py — Gera briefing de aprovação .docx a partir de Markdown.

Pipeline:
1. Lê briefing-aprovacao.md (gerado pela skill)
2. Renderiza .docx formatado em padrão editorial
3. Salva em 03-revisoes/Briefing-Aprovacao-{SIGLA}-{DDMMYY}.docx

Uso:
    python3 gerar_briefing_docx.py \
        --md _private/briefings/briefing-aprovacao-mpr.md \
        --out 03-revisoes/Briefing-Aprovacao-MPR-120526.docx \
        --sigla MPR \
        --expert "Marina Costa"

Dependências:
    pip install python-docx
"""
import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERRO: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# Paleta editorial (Editorial Premium do DESIGN.md)
INK = RGBColor(0x0A, 0x0A, 0x0A)
PAPER = RGBColor(0xF4, 0xF1, 0xEA)
GOLD = RGBColor(0xD4, 0xAF, 0x37)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_inline_runs(paragraph, text):
    """Suporta **bold**, *italic*, `code`."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def configure_styles(doc):
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    s.font.color.rgb = INK
    for level, size in [("Heading 1", 28), ("Heading 2", 20), ("Heading 3", 14), ("Heading 4", 12)]:
        st = doc.styles[level]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = INK


def add_cover(doc, sigla, nome_evento, big_idea, expert, nicho, data_inicio, edicao):
    # Linha topo: edição
    p = doc.add_paragraph()
    r = p.add_run(f"EDIÇÃO {edicao or '001'} · {sigla}")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    r.font.bold = True

    # Espaço
    for _ in range(2):
        doc.add_paragraph()

    # Título do evento
    p = doc.add_paragraph()
    r = p.add_run(nome_evento)
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = INK

    # Big idea (italic em serif)
    p = doc.add_paragraph()
    r = p.add_run(big_idea)
    r.italic = True
    r.font.size = Pt(18)
    r.font.color.rgb = MUTED

    # Espaço
    for _ in range(4):
        doc.add_paragraph()

    add_horizontal_rule(doc)

    # Metadata
    meta = [
        ("Especialista", expert or "—"),
        ("Nicho", nicho or "—"),
        ("Data de início", data_inicio or "—"),
        ("Status", "🔴 Aguardando aprovação"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.autofit = False
    for ri, (k, v) in enumerate(meta):
        c0 = tbl.rows[ri].cells[0]
        c1 = tbl.rows[ri].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(k)
        r0.bold = True
        c1.paragraphs[0].add_run(v)
        c0.width = Inches(2)
        c1.width = Inches(4.5)

    add_page_break(doc)


def render_markdown_simplified(doc, md_text):
    """Renderiza Markdown → .docx · suporta H1-H4, listas, blockquotes, tabelas, code."""
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        # Headers
        if s.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_runs(p, s[2:].strip())
            i += 1
            continue
        if s.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, s[3:].strip())
            i += 1
            continue
        if s.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_runs(p, s[4:].strip())
            i += 1
            continue
        if s.startswith("#### "):
            p = doc.add_paragraph()
            r = p.add_run(s[5:].strip())
            r.bold = True
            r.font.size = Pt(12)
            i += 1
            continue

        # Horizontal rule
        if s in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Code fence
        if s.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            i += 1
            continue

        # Tables (pipe syntax)
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = tbl.rows[0].cells
            for ci, h in enumerate(header):
                hdr[ci].text = ""
                p = hdr[ci].paragraphs[0]
                r = p.add_run(h)
                r.bold = True
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    if ci < len(tbl.rows[ri + 1].cells):
                        tbl.rows[ri + 1].cells[ci].text = ""
                        add_inline_runs(tbl.rows[ri + 1].cells[ci].paragraphs[0], val)
            doc.add_paragraph()
            continue

        # Lists
        if re.match(r"^\s*[-*+]\s+", line):
            while i < len(lines) and re.match(r"^\s*[-*+]\s+", lines[i]):
                m = re.match(r"^(\s*)[-*+]\s+(.*)", lines[i])
                p = doc.add_paragraph(style="List Bullet")
                add_inline_runs(p, m.group(2))
                i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                m = re.match(r"^\s*\d+\.\s+(.*)", lines[i])
                p = doc.add_paragraph(style="List Number")
                add_inline_runs(p, m.group(1))
                i += 1
            continue

        # Blockquote
        if s.startswith(">"):
            quote = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip()[1:].strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(" ".join(quote))
            r.italic = True
            r.font.color.rgb = MUTED
            r.font.size = Pt(11)
            continue

        # Empty
        if not s:
            i += 1
            continue

        # Paragraph
        para_lines = [s]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^\s*(#{1,6}\s|[-*+]\s|\d+\.\s|>|```|---|\*\*\*|___)", lines[i]
        ) and "|" not in lines[i]:
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(para_lines))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--md", required=True, help="Path do briefing-aprovacao.md")
    ap.add_argument("--out", required=True, help="Path destino .docx")
    ap.add_argument("--sigla", required=True, help="Sigla do projeto (3 letras maiúsculas)")
    ap.add_argument("--expert", required=True, help="Nome do especialista")
    ap.add_argument("--nicho", default="", help="Nicho do expert")
    ap.add_argument("--evento", default="", help="Nome do evento")
    ap.add_argument("--big-idea", default="", help="Big idea (1 frase)")
    ap.add_argument("--data-inicio", default="", help="Data início (DD/MM/AAAA)")
    ap.add_argument("--edicao", default="001", help="Número da edição")
    args = ap.parse_args()

    md_path = Path(args.md)
    if not md_path.exists():
        print(f"ERRO: {md_path} não existe", file=sys.stderr)
        sys.exit(1)

    md_text = md_path.read_text(encoding="utf-8")

    # Build doc
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11.69)  # A4
    sec.page_width = Inches(8.27)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    configure_styles(doc)
    add_cover(
        doc,
        sigla=args.sigla,
        nome_evento=args.evento or "Briefing de Aprovação",
        big_idea=args.big_idea or "Aguardando big idea",
        expert=args.expert,
        nicho=args.nicho,
        data_inicio=args.data_inicio,
        edicao=args.edicao,
    )

    render_markdown_simplified(doc, md_text)

    # Salva
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"OK · {out_path.absolute()}")


if __name__ == "__main__":
    main()
