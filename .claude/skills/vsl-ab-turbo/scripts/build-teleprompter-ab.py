#!/usr/bin/env python3
# ─────────────────────────────────────────────────────────────────────────────
# build-teleprompter-ab.py · gerador do teleprompter A/B (skill vsl-ab-turbo)
#
# Gera o .docx de teleprompter pra GRAVAÇÃO ÚNICA de duas versões de VSL:
#   Parte 1 · Versão A completa (hook ao fechamento — takes compartilhados saem daqui)
#   Parte 2 · só os takes divergentes da Versão B
#   Parte 3 · hooks alternativos (opcional)
#
# USO:
#   1. Copie este arquivo pra pasta do projeto.
#   2. Preencha TITULO, SUBTITULO, ARQUIVO e a lista BLOCOS com o roteiro
#      JÁ REVISADO pelo @revisor-copy-turbo (só o texto FALADO).
#      - IDs de take conforme references/mapa-de-takes.md (A1, B1, 2C, 4A, EXTRA...)
#      - Direção padrão entre takes: a constante P3 (ponto de corte do editor)
#      - Números que travam a leitura: por extenso ("998 reais")
#   3. Rode num venv (python-docx não instala no Python do sistema · PEP 668):
#        python3 -m venv /tmp/vsl-venv && /tmp/vsl-venv/bin/pip install -q python-docx
#        /tmp/vsl-venv/bin/python build-teleprompter-ab.py
#
# Formato: só o falado · Arial 30pt · marcadores discretos em dourado ·
# direções em cinza pequeno "não fale".
# ─────────────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── PREENCHER ────────────────────────────────────────────────────────────────
TITULO    = "VSL · [PRODUTO] — [EXPERT]"
SUBTITULO = "Teleprompter · gravação única das versões A e B · revisado (@revisor-copy-turbo) · alvo XX min por versão"
ARQUIVO   = "VSL-[PRODUTO]-Teleprompter.docx"

P3 = "PAUSA DE 3 SEGUNDOS · olha pra câmera · retoma — ponto de corte do editor"

# Cada bloco: ("MARCADOR DO TAKE", [parágrafos falados], "direção (não lida) ou None")
# Cabeçalhos de parte: lista de parágrafos vazia + direção explicando a parte.
BLOCOS = [
    ("PARTE 1 — VERSÃO A COMPLETA · gravar direto, do take A1 ao fechamento", [],
     "esta parte vira a VSL A inteira; a VSL B reaproveita os takes compartilhados daqui"),

    ("TAKE A1 · HOOK — VERSÃO A (dor ampla)", [
        "[hook A — sem nomear o método]",
        "[continuação curta do hook]",
    ], P3),

    ("TAKE A2 · [BLOCO 2] — abertura VERSÃO A", [
        "[abertura A do problema]",
    ], P3),

    ("TAKE 2C · [BLOCO 2] — continuação (COMPARTILHADA · entra nas duas versões)", [
        "[continuação do problema, idêntica nas duas versões]",
    ], P3),

    ("TAKE A3 · [BLOCO 3] — VERSÃO A (sem falar o nome do método)", [
        "[ponte de credibilidade, versão A]",
    ], P3),

    ("TAKE 4 · MECANISMO — parte 1 (COMPARTILHADA)", [
        "[mecanismo até o ponto da revelação]",
    ], P3),

    ("TAKE 4A · REVELAÇÃO DO MÉTODO — SÓ ENTRA NA VERSÃO A (pausa antes e depois)", [
        "[parágrafo que nomeia o método pela primeira vez]",
    ], P3),

    ("TAKE 4F · MECANISMO — fechamento (COMPARTILHADO)", [
        "[fechamento do mecanismo]",
    ], P3),

    # ... TAKES 5-11 compartilhados (quebra de crença, prova, depoimentos,
    #     future pacing, oferta, FAQ em mini-takes, escassez, fechamento) ...

    ("PARTE 2 — SÓ OS TRECHOS DA VERSÃO B (takes curtos)", [],
     "o editor monta a VSL B com estes takes + os compartilhados da Parte 1"),

    ("TAKE B1 · HOOK — VERSÃO B (método explícito)", [
        "[hook B — nomeia o método]",
    ], P3),

    # ... B2, B3 ...

    ("PARTE 3 — HOOKS ALTERNATIVOS (opcional · viram aberturas de teste e cortes de anúncio)", [],
     "gravar os que tiver energia — cada um é um take independente de ~15s"),

    ("EXTRA A2 · [ângulo]", [
        "[hook alternativo]",
    ], None),
]
# ─────────────────────────────────────────────────────────────────────────────

GOLD = RGBColor(0xC9, 0xA2, 0x27)
GRAY = RGBColor(0x88, 0x88, 0x88)
INK  = RGBColor(0x10, 0x10, 0x10)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.6)
    s.left_margin = s.right_margin = Inches(0.7)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(30)
normal.font.color.rgb = INK
normal.paragraph_format.line_spacing = 1.4
normal.paragraph_format.space_after = Pt(14)

def add(text, size, color, bold=False, italic=False, after=14, before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color
    r.bold = bold; r.italic = italic; r.font.name = "Arial"
    return p

add(TITULO, 16, GOLD, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
add(SUBTITULO, 11, GRAY, italic=True, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
add("Leia só o texto grande. Linhas em cinza pequeno = orientação, não fale.",
    11, GRAY, italic=True, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)

for marcador, paragrafos, direcao in BLOCOS:
    add("▌ " + marcador, 15, GOLD, bold=True, before=18, after=8)
    for i, txt in enumerate(paragrafos):
        last = (i == len(paragrafos) - 1) and direcao is None
        add(txt, 30, INK, after=(20 if last else 14))
    if direcao:
        add("( " + direcao + " )", 13, GRAY, italic=True, after=20)

doc.save(ARQUIVO)
print(f"✓ {ARQUIVO} gerado")
