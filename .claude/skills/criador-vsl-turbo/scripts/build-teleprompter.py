#!/usr/bin/env python3
# ─────────────────────────────────────────────────────────────────────────────
# build-teleprompter.py · gerador do .docx pronto pra teleprompter (skill criador-vsl-turbo)
#
# USO:
#   1. Copie este arquivo pra pasta do projeto (02-entregaveis-finais/vsl/).
#   2. Preencha TITULO, SUBTITULO e a lista BLOCOS com o roteiro JÁ REVISADO
#      pelo @revisor-copy-turbo (só o texto FALADO).
#   3. Rode num venv (python-docx não roda no Python do sistema · PEP 668):
#        python3 -m venv /tmp/vsl-venv && /tmp/vsl-venv/bin/pip install -q python-docx
#        /tmp/vsl-venv/bin/python build-teleprompter.py
#
# Formato (ver references/saida-teleprompter.md): só o falado · Arial 30pt ·
# marcadores discretos em dourado · direções em cinza pequeno "não fale".
# ─────────────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── PREENCHER ────────────────────────────────────────────────────────────────
TITULO    = "VSL · [NOME DO PROJETO] — [EXPERT]"
SUBTITULO = "Roteiro de teleprompter · versão revisada (@revisor-copy-turbo) · ~XX min"
ARQUIVO   = "VSL-[PROJETO]-Teleprompter.docx"

# Cada bloco: ("MARCADOR", [parágrafos falados], "direção (não lida) ou None")
BLOCOS = [
    ("1 · ABERTURA", [
        "[hook · primeira frase que para o scroll · abre o loop]",
        "[continuação curta do hook]",
    ], None),
    ("2 · PROBLEMA", [
        "[problema + agitação · estado-atual sensorial]",
    ], None),
    ("3 · QUEM SOU EU", [
        "[ponte de credibilidade · herói relutante]",
    ], None),
    ("4 · O MECANISMO", [
        "[UMP — causa-raiz oculta nomeada]",
        "[UMS — por que sua solução resolve · fecha o loop do hook]",
    ], None),
    ("5 · QUEBRA DE CRENÇA", [
        "[escada de crenças · remove culpa · 'não é sua culpa, faltava X']",
    ], None),
    ("6 · PROVA", [
        "[caso narrativo com número/nome + imperfeição]",
    ], "prints/depoimentos reais aqui, autorizados, marcados “resultado individual, pode variar”"),
    ("7 · COMO VAI SER", [
        "[future pacing · estado-céu · presente · sensorial]",
    ], None),
    ("8 · A OFERTA", [
        "[reveal → stack → ancoragem → preço → garantia]",
    ], None),
    ("9 · DÚVIDAS", [
        "[FAQ · objeções finais: tempo, funciona pra mim, preço]",
    ], None),
    ("10 · AGORA", [
        "[escassez/urgência — preencher data ou nº de vagas REAL]",
    ], "preencher escassez REAL antes de gravar — sem verdade, reprova na Meta"),
    ("11 · FECHAMENTO", [
        "[recap → reafirma mecanismo → CTA único → o que acontece ao clicar]",
    ], None),
]
# ─────────────────────────────────────────────────────────────────────────────

GOLD = RGBColor(0xC9, 0xA2, 0x27)
GRAY = RGBColor(0x88, 0x88, 0x88)
INK  = RGBColor(0x10, 0x10, 0x10)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.6)
    s.left_margin = s.right_margin = Inches(0.7)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(30)
normal.font.color.rgb = INK
normal.paragraph_format.line_spacing = 1.4
normal.paragraph_format.space_after = Pt(14)

def add(text, size, color, bold=False, italic=False, after=14, before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color
    r.bold = bold; r.italic = italic; r.font.name = "Arial"
    return p

add(TITULO, 16, GOLD, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
add(SUBTITULO, 11, GRAY, italic=True, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
add("Leia só o texto grande. Linhas em cinza pequeno = orientação, não fale.",
    11, GRAY, italic=True, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)

for marcador, paragrafos, direcao in BLOCOS:
    add("▌ " + marcador, 15, GOLD, bold=True, before=18, after=8)
    for i, txt in enumerate(paragrafos):
        last = (i == len(paragrafos) - 1) and direcao is None
        add(txt, 30, INK, after=(20 if last else 14))
    if direcao:
        add("( " + direcao + " )", 13, GRAY, italic=True, after=20)

doc.save(ARQUIVO)
print(f"✓ {ARQUIVO} gerado")
